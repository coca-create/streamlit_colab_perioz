
import os
import re
import tempfile
import pandas as pd
import codecs
from datetime import datetime
from docx import Document
from streamlit_module import moz_tab7 as t7
from streamlit_module import moz_tab3 as t3
from streamlit_module import moz_tab1 as t1

'''def unify_timestamps_vtt(text):
    # Define patterns for different timestamp formats
    pattern_1_digit = re.compile(r'(\d{2}:\d{2}:\d{2}\.\d)(?!\d)')
    pattern_2_digits = re.compile(r'(\d{2}:\d{2}:\d{2}\.\d{2})(?!\d)')
    pattern_3_digit = re.compile(r'(\d{1}:\d{2}:\d{2}\.\d)(?!\d)')
    pattern_4_digits = re.compile(r'(\d{1}:\d{2}:\d{2}\.\d{2})(?!\d)')

    # Replace 1-digit and 2-digit millisecond formats with 3-digit format for HH:MM:SS format
    if pattern_1_digit.search(text):
        text = pattern_1_digit.sub(lambda x: x.group(1) + '00', text)
        #print("pattern1")
    if pattern_2_digits.search(text):
        text = pattern_2_digits.sub(lambda x: x.group(1) + '0', text)
        #print("pattern2")

    # Replace 1-digit and 2-digit millisecond formats with 3-digit format for H:MM:SS format
    if pattern_3_digit.search(text):
        text = pattern_3_digit.sub(lambda x: x.group(1) + '00', text)
        #print("pattern3")
    if pattern_4_digits.search(text):
        text = pattern_4_digits.sub(lambda x: x.group(1) + '0', text)
        #print("pattern4")

    return text


#srtファイルのタイムスタンプ桁数を統一。
def unify_timestamps(text):
    # Define patterns for different timestamp formats
    pattern_1_digit = re.compile(r'(\d{2}:\d{2}:\d{2}\,\d)(?!\d)')
    pattern_2_digits = re.compile(r'(\d{2}:\d{2}:\d{2}\,\d{2})(?!\d)')
    
    # Replace 1-digit millisecond format with 3-digit format
    text = pattern_1_digit.sub(lambda x: x.group(1) + '00', text)
    
    # Replace 2-digit millisecond format with 3-digit format
    text = pattern_2_digits.sub(lambda x: x.group(1) + '0', text)
    
    return text'''

def unify_timestamps_vtt(text):
    # HH:MM:SSとH:MM:SSのタイムスタンプ形式
    pattern_1_digit_hh = re.compile(r'(\d{2}:\d{2}:\d{2}\.\d)(?!\d)')  # 2桁の時間
    pattern_2_digits_hh = re.compile(r'(\d{2}:\d{2}:\d{2}\.\d{2})(?!\d)')
    
    pattern_1_digit_h = re.compile(r'(\d{1}:\d{2}:\d{2}\.\d)(?!\d)')   # 1桁の時間
    pattern_2_digits_h = re.compile(r'(\d{1}:\d{2}:\d{2}\.\d{2})(?!\d)')

    # HH:MM:SS形式とH:MM:SS形式のミリ秒を統一
    text = pattern_1_digit_hh.sub(lambda x: x.group(1) + '00', text)
    text = pattern_2_digits_hh.sub(lambda x: x.group(1) + '0', text)
    text = pattern_1_digit_h.sub(lambda x: x.group(1) + '00', text)
    text = pattern_2_digits_h.sub(lambda x: x.group(1) + '0', text)

    return text

# SRT用のタイムスタンプを統一する関数
def unify_timestamps(text):
    # HH:MM:SSとH:MM:SSのタイムスタンプ形式
    pattern_1_digit_hh = re.compile(r'(\d{2}:\d{2}:\d{2}\,\d)(?!\d)')  # 2桁の時間
    pattern_2_digits_hh = re.compile(r'(\d{2}:\d{2}:\d{2}\,\d{2})(?!\d)')
    
    pattern_1_digit_h = re.compile(r'(\d{1}:\d{2}:\d{2}\,\d)(?!\d)')   # 1桁の時間
    pattern_2_digits_h = re.compile(r'(\d{1}:\d{2}:\d{2}\,\d{2})(?!\d)')

    # HH:MM:SS形式とH:MM:SS形式のミリ秒を統一
    text = pattern_1_digit_hh.sub(lambda x: x.group(1) + '00', text)
    text = pattern_2_digits_hh.sub(lambda x: x.group(1) + '0', text)
    text = pattern_1_digit_h.sub(lambda x: x.group(1) + '00', text)
    text = pattern_2_digits_h.sub(lambda x: x.group(1) + '0', text)

    return text


def unify_timestamps_forlist(lines, format_type):
    if format_type == 'vtt':
        unify_timestamps_func = unify_timestamps_vtt
    elif format_type == 'srt':
        unify_timestamps_func = unify_timestamps
    else:
        raise ValueError(f"Unsupported format type: {format_type}")

    return [unify_timestamps_func(line) if '-->' in line else line for line in lines]

def read_file_content(file):
    if file is None:
        return """<div style='color: orange !important; font-family: inherit; text-align: center; 
                display: flex; align-items: flex-start; justify-content: center; height: 400px; padding-top: 40px;'>
                No file uploaded
                </div>"""

    file_extension = os.path.splitext(file)[1]
    content = ""

    if file_extension == '.txt':#file.nameの.name除去中
        with codecs.open(file, 'r', 'utf-8') as f:
            content = f.read()
            content = f"""<pre style="white-space: pre-wrap; overflow-y: auto; height: 500px; word-wrap: break-word; padding: 10px; font-family: inherit; font-size: inherit;">{content}</pre>"""

    elif file_extension == '.srt' : # file.nameのname除去中。
        with codecs.open(file, 'r', 'utf-8') as f:
            content = f.read()
            content=unify_timestamps(content)
            content = f"""<pre style="white-space: pre-wrap; overflow-y: auto; height: 500px; word-wrap: break-word; padding: 10px; font-family: inherit; font-size: inherit;">{content}</pre>"""

    elif file_extension == '.vtt': # file.nameのname除去中。
        with codecs.open(file,'r',encoding='utf-8')as f:
            content = f.read()
            content = unify_timestamps_vtt(content)
            content = f"""<pre style="white-space: pre-wrap; overflow-y: auto; height: 500px; word-wrap: break-word; padding: 10px; font-family: inherit; font-size: inherit;">{content}</pre>"""
    else:
        return None
    return content

def display_file_content(file):
    if file is None:
        return read_file_content(file),[],pd.DataFrame({'1': [''], '2': [''], '3': ['']})
    if file.endswith('.txt'):
        filename = os.path.basename(file)    
        match = re.match(r"(.+?)(_NR\.txt|_R\.txt)$", filename)
        
        if match:
            basename, ext = match.groups()
            if ext == "_NR.txt":

                timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
                temp_dir = os.path.join(tempfile.gettempdir(), f"tempdir_{timestamp}")
                os.makedirs(temp_dir, exist_ok=True)

                doc_filename = os.path.join(temp_dir, f"{basename}_txtnr.docx")
            elif ext == "_R.txt":
                timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
                temp_dir = os.path.join(tempfile.gettempdir(), f"tempdir_{timestamp}")
                os.makedirs(temp_dir, exist_ok=True)
                doc_filename = os.path.join(temp_dir, f"{basename}_txtr.docx")

            with open(file, 'r', encoding='utf-8') as f:
                content = f.read()
            doc = Document()
            doc.add_paragraph(content)
            doc.save(doc_filename)

            return read_file_content(file), doc_filename,''
              
    else:
        if file.endswith('.srt'):
            excel_path,t4_df = t3.create_excel_from_srt(english_path=file,tail="")
        elif file.endswith('.vtt'):
            excel_path,t4_df = t3.create_excel_from_srt(english_path=file,tail="")

        t4_df=t1.dataframe(t4_df)

    return read_file_content(file), excel_path,t4_df

def save_translated_content(file, translated_text):
    if file==None:
        return []
    
    file_name, file_extension = os.path.splitext(file) 
    output_file_path = file_name + "_ja" + file_extension

    '''    if file_extension == '.docx':
            doc = Document()
            doc.add_paragraph(translated_text)
            doc.save(output_file_path)
    '''
    if file_extension == '.txt':
        with codecs.open(output_file_path, 'w', 'utf-8') as f:
            f.write(translated_text)
    
    elif  file_extension == '.srt' or file_extension == '.vtt':
        content = re.sub(r'[\u200B-\u200D\uFEFF]', '', translated_text)
        content = content.replace('\u200B', '')

        if "-->" in content:
            content = re.sub(r'\s+', '', content)

            if file_extension == '.srt':
                pattern = re.compile(r'(\d{1,4})(\d{2}:\d{2}:\d{2},\d{3}-->\d{2}:\d{2}:\d{2},\d{3})')
            elif file_extension == '.vtt':
                content=t7.webvtt_remover(content)
                pattern = re.compile(r'(\d{1,4})(\d{1}:\d{2}:\d{2}.\d{3}-->\d{1}:\d{2}:\d{2}.\d{3})')
            matches = pattern.findall(content)      
            #print(f"d1(len(matches)):{len(matches)}")  
            if len(matches)==0:
                pattern = re.compile(r'(\d{1,4})(\d{2}:\d{2}:\d{2}.\d{3}-->\d{2}:\d{2}:\d{2}.\d{3})')
                matches = pattern.findall(content)
                #print(f"d2(len(matches)):{len(matches)}")  
            if len(matches)==0:
                pattern = re.compile(r'(\d{1,4})(\d{1}:\d{2}:\d{2},\d{3}-->\d{1}:\d{2}:\d{2},\d{3})')
                matches = pattern.findall(content)
            segments = pattern.split(content)
            corrected_content = []
            
            for i in range(1, len(segments), 3):
                segment_id = segments[i]
                timestamp = segments[i + 1]
                text = segments[i + 2]
                
                corrected_content.append(f"{segment_id}")
                corrected_content.append(timestamp.replace('-->', ' --> '))
                corrected_content.append(text)


        else:

            # Normalize line endings
            content = re.sub(r'\r\n', '\n', content)
            # Remove multiple empty lines
            content = re.sub(r'\n\n+', '\n', content)
            # Normalize spaces
            content = re.sub(r'\t+', '\t', content)

            #print(f"After normalization: {repr(content)}")  # デバッグ用に表示

            # Split content into lines and process each line
            content = content.strip().split('\n')
            corrected_content = []

            for line in content:
                parts = line.split('\t')
                if len(parts) < 4:
                    continue  # 不正な行をスキップ
                index = parts[0].strip()
                start_time = parts[1].strip()
                end_time = parts[2].strip()
                text = parts[3].strip()
                
                corrected_content.append(index)
                corrected_content.append(f"{start_time} --> {end_time}")
                corrected_content.append(text)
                corrected_content.append('')

            # Remove any trailing empty elements
            while corrected_content and corrected_content[-1] == '':
                corrected_content.pop()

            # Ensure each subtitle block is separated by exactly one empty line
            if file_extension == '.vtt':
                #corrected_content = "WEBVTT\n\n" + "\n\n".join("\n".join(corrected_content[i:i+3]) for i in range(0, len(corrected_content), 3))                
                corrected_content = "WEBVTT\n\n" + "\n\n".join("\n".join(corrected_content[i:i+4]) for i in range(0, len(corrected_content), 4))#結局４つじゃないと駄目。
            else:
                #corrected_content = "\n\n".join("\n".join(corrected_content[i:i+3]) for i in range(0, len(corrected_content), 3))
                corrected_content = "\n\n".join("\n".join(corrected_content[i:i+4]) for i in range(0, len(corrected_content), 4))

            with open(output_file_path, 'w', encoding='utf-8') as file:
                file.write(corrected_content + '\n')
              
    return output_file_path

def translate(file, translated_text):
    return save_translated_content(file, translated_text)
