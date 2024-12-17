import streamlit as st
import os
import re
import codecs
from docx import Document
import pandas as pd
from openpyxl.styles import Alignment, Font, PatternFill
import streamlit.components.v1 as components
import zipfile
import tempfile
from datetime import datetime
from streamlit_module import moz_tab5 as t5
from streamlit_module import moz_tab4 as t4
from streamlit_module import moz_tab3 as t3
from streamlit_module import moz_tab7 as t7
from streamlit_module import moz_tab8 as t8
from streamlit_module import moz_split as cm
from streamlit_module import common as co
from streamlit_module import moz_split_period as period
import csv
#from streamlit_module import moz_tab1 as t1


# セッション状態をクリアする関数
def clear_session_state(tab=None,select_clear=None,option=None):#selsect_clearはtab内に複数のクリアボタン,optionは1つのクリアボタンに複数のradiobuttonケースを対応させる.

    tab1_keys = ['t1_input_file_path','t1_uploaded_file','t1_translated_text','output_html_tab1', 'output_files_tab1', 'output_file_tab1', 'output_ja_files_tab1', 'df_data_tab1']
    tab2_keys = ['df_data_tab2', 't2_translated_text','file_path_tab2','df_tab2','output_html_tab2', 'input_file_path_tab2','output_file_tab2', 'translated_text_tab2',  'saved_files_tab2']
    tab3_keys = ['t3_namelist_B','t3_namelist_A','output_files_tab3_docx_to_srttxt', 'output_files_tab3_srttxt_to_docx','t3_upload_files_A','t3_upload_files_B']
    #tab4_keys = ['output_excel_tab4', 'df_data_tab4', 'download_filename_tab4','english_path_both','japanese_path_both','english_path_only','japanese_path_only']
    tab4_keys_both=['english_path_both','japanese_path_both','both_excel_file_path','both_df','both_download_filename_tab4']
    tab4_keys_en=['english_path_only','en_excel_file_path','en_df','en_download_filename_tab4']
    tab4_keys_ja=['japanese_path_only','ja_excel_file_path','ja_df','en_download_filename_tab4']
    tab5_keys = ['upload_files','t5_kuten_files','t5_filenames', 'results']
    tab6_keys = ['t6_srtnamelist','t6_jsnamelist','t6_jsons','t6_srts','tab9_output_srt_files','tab9_output_txt_files','tab9_files_ready']#'tab9_output_srt_file','tab9_output_txt_file',
    tab7_keys = ['yt_namelist','yt_output_srt','yt_output_txt','download_ready','yt_srts']
    tab11_keys = ['t11_namelist','t11_srt_files','t11_jsons','skipped_files']

    if tab == '配布字幕の再編':
        for key in tab1_keys:
            if key in st.session_state:
                del st.session_state[key]
            st.session_state[key] = None  # Reset value to None
        st.session_state.uploader_key += 1    
        
        
       
    elif tab == '翻訳お手伝い':
        for key in tab2_keys:
            if key in st.session_state:
                del st.session_state[key]
            st.session_state[key] = None  # Reset value to None
        st.session_state.uploader_tab2_key += 1  

    elif tab == 'Word/Excel↔SRT/VTT/TXT':
        if select_clear=="First":
            tab3_keys=['t3_namelist_A','output_files_tab3_docx_to_srttxt','t3_upload_files_A']
            st.session_state.uploader_tab3_1_key += 1
        if select_clear=="Second":
            tab3_keys=['t3_namelist_B','output_files_tab3_srttxt_to_docx','t3_upload_files_B']
            st.session_state.uploader_tab3_2_key += 1
        else:
            st.session_state.uploader_tab3_1_key += 1
            st.session_state.uploader_tab3_2_key += 1

        for key in tab3_keys:
            if key in st.session_state:
                del st.session_state[key]
            st.session_state[key] = []  # Reset value to empty list
        

    elif tab == 'SRT/VTT→Excel(2言語)':
      
        if option=="日本語と英語":
            st.session_state.uploader_tab4_1_key += 1
            st.session_state.uploader_tab4_2_key += 1
            for key in tab4_keys_both:
                if key in st.session_state:            
                    del st.session_state[key]
                st.session_state[key]=None

        elif option=="英語のみ":
            st.session_state.uploader_tab4_3_key += 1
            for key in tab4_keys_en:
                if key in st.session_state:            
                  del st.session_state[key]
                st.session_state[key]=None

        elif option=="日本語のみ":
            st.session_state.uploader_tab4_4_key += 1
            for key in tab4_keys_ja:
                if key in st.session_state:            
                    del st.session_state[key]
                st.session_state[key]=None
        else:
            st.session_state.uploader_tab4_1_key += 1
            st.session_state.uploader_tab4_2_key += 1
            st.session_state.uploader_tab4_3_key += 1
            st.session_state.uploader_tab4_4_key += 1
            for key in tab4_keys_both+tab4_keys_en+tab4_keys_ja:
                if key in st.session_state:            
                    del st.session_state[key]
                st.session_state[key]=None

        
        

    elif tab == '日本語srt,vttの句読点分割':
        for key in tab5_keys:
            if key in st.session_state:
                del st.session_state[key]
            st.session_state[key] = None
        st.session_state.uploader_tab5_key += 1
            
    elif tab == 'whisperファイルの復活①':
        for key in tab6_keys:
            if key in st.session_state:   
                del st.session_state[key]
            st.session_state[key] = ""
        st.session_state['tab9_files_ready'] = False
        st.session_state.uploader_tab6_1_key += 1
        st.session_state.uploader_tab6_2_key += 1

    elif tab == 'whisperファイルの復活②':
        for key in tab11_keys:
            if key in st.session_state:
                del st.session_state[key]
            st.session_state[key]=None
        st.session_state.uploader_tab7_key += 1

    elif tab == 'YT付属字幕の再編':
        for key in tab7_keys:
            if key in st.session_state:
                del st.session_state[key]
            st.session_state[key]=None
        st.session_state['download_ready']=False
        st.session_state.uploader_tab8_key +=1
    st.session_state.current_tab = tab
    st.rerun()

def restore_special_periods(text):
    text = text.replace('[dot]', '.')
    return text



def split_segment(segment, start_time, end_time): # 文字数基準の分割。
    if start_time is None or end_time is None:
        return [(segment, start_time, end_time)]

    #print(f"Segment before split: {segment}")  # デバッグ出力

    # 空白で区切って単語単位で処理
    words = segment.split()

    # 分割用リスト
    sentences = []
    current_sentence = []

    for word in words:
        current_sentence.append(word)
        # 単語がピリオド, 感嘆符, 質問符で終わる場合
        if word.endswith(('.', '!', '?')):
            # 現在の文を文字列として結合し、リストに追加
            sentences.append(' '.join(current_sentence))
            current_sentence = []

    # 最後の文があれば追加
    if current_sentence:
        sentences.append(' '.join(current_sentence))

    # 文の長さ（文字数）を考慮して、時間を配分
    total_length = sum(len(sentence) for sentence in sentences)
    if total_length == 0:
        return [(segment, start_time, end_time)]  # 文がない場合、元のセグメントをそのまま返す
    
    #print(f"Sentences after split: {sentences}")  # デバッグ出力

    # タイムスタンプを文の長さに応じて割り振る
    times = [start_time]
    for sentence in sentences:
        sentence_length = len(sentence)
        duration = (end_time - start_time) * (sentence_length / total_length)
        times.append(times[-1] + duration)

    return [(sentences[i], times[i], times[i+1]) for i in range(len(sentences))]

#単語数基準で分割したいなら、↓




def merge_segments(segments):
    merged_segments = []
    buffer_segment = ""
    buffer_start = None
    buffer_end = None

    for segment, start, end in segments:
        if buffer_segment:
            buffer_segment += " " + segment
            buffer_end = end
            if segment.endswith('.') or segment.endswith('?') or segment.endswith('!'):
                merged_segments.append((buffer_segment, buffer_start, buffer_end))
                buffer_segment = ""
                buffer_start = None
                buffer_end = None
        else:
            if segment.endswith('.') or segment.endswith('?') or segment.endswith('!'):
                merged_segments.append((segment, start, end))
            else:
                buffer_segment = segment
                buffer_start = start
                buffer_end = end

    if buffer_segment:
        merged_segments.append((buffer_segment, buffer_start, buffer_end))

    #print(f"Merged Segments: {merged_segments}")  # デバッグ出力

    return merged_segments
def process_vtt(lines,replace_words):
    segments = []
    start_time = None
    end_time = None
    text = ""
    header = lines[0]

    for line in lines[1:]:
        if re.match(r'^\d+$', line.strip()):
            if text:               
                common_app_data = os.getenv('APPDATA')  # WindowsのCommon AppDataフォルダ
                data_folder = "/content/my_app"
                if st.session_state.select_dot_result:
                    with open(os.path.join(data_folder,st.session_state.select_dot_result), newline='',encoding='utf-8') as dot_csvfile:
                        reader = csv.reader(dot_csvfile)
                        next(reader)
                        dot_replacements = [(row[0],row[1]) for row in reader]
                    
                    
                    for dot_original, dot_replacement in dot_replacements:
                        r_dot_original=re.escape(dot_original)
                        dot_new_original = rf"\b{r_dot_original}"
                        text = re.sub(dot_new_original, dot_replacement, text)

                if start_time is not None and end_time is not None:
                    segments.extend(split_segment(text.strip(), start_time, end_time))
            text = ""
            
        elif '-->' in line:
            times = line.strip().split(' --> ')
            start_time = convert_time_to_seconds(times[0])
            end_time = convert_time_to_seconds(times[1])
        else:
            text += line.strip() + " "

    if text:
        #print(st.session_state.select_dot_result)
        if st.session_state.select_dot_result:
            common_app_data = os.getenv('APPDATA')  # WindowsのCommon AppDataフォルダ
            data_folder = "/content/my_app"
            
            with open(os.path.join(data_folder,st.session_state.select_dot_result), newline='',encoding='utf-8') as dot_csvfile:
                reader = csv.reader(dot_csvfile)
                next(reader)
                dot_replacements = [(row[0],row[1]) for row in reader]
            
            
            for dot_original, dot_replacement in dot_replacements:
                r_dot_original=re.escape(dot_original)
                dot_new_original = rf"\b{r_dot_original}"
                text = re.sub(dot_new_original, dot_replacement, text)

        #text = replace_special_periods(text)
        if start_time is not None and end_time is not None:
            segments.extend(split_segment(text.strip(), start_time, end_time))

    merged_segments = merge_segments(segments)


    if replace_words==True and st.session_state.select_rp_result: 
        common_app_data = os.getenv('APPDATA')  # WindowsのCommon AppDataフォルダ
        data_folder = "/content/my_app"
        with open(os.path.join(data_folder,st.session_state.select_rp_result), newline='', encoding='utf-8') as csvfile:
            reader = csv.reader(csvfile)
            next(reader)  # ヘッダーをスキップ
            replacements = [(row[0], row[1]) for row in reader]


    output = [header]
    segment_number = 0
    for text, start, end in merged_segments:        
        text = restore_special_periods(text)
        if replace_words==True and st.session_state.select_rp_result:
            for original, replacement in replacements:
                original=re.escape(original)
                new_original = rf"\b{original}\b"
                text = re.sub(new_original, replacement, text)
        output.append(f"{segment_number + 1}")
        output.append(f"{convert_seconds_to_time(start, 'vtt')} --> {convert_seconds_to_time(end, 'vtt')}")
        output.append(text)
        segment_number += 1
        output.append("")  # セグメント間の改行を保持

    return '\n'.join(output)

def process_srt(lines,replace_word):
    segments = []
    start_time = None
    end_time = None
    text = ""
    segment_index = 0

    for line in lines:
        if re.match(r'^\d+$', line.strip()):
            if text:
                # text = replace_special_periods(text)
                if st.session_state.select_dot_result:
                    common_app_data = os.getenv('APPDATA')  # WindowsのCommon AppDataフォルダ
                    data_folder = "/content/my_app"
                    with open(os.path.join(data_folder,st.session_state.select_dot_result), newline='',encoding='utf-8') as dot_csvfile:
                        reader = csv.reader(dot_csvfile)
                        next(reader)
                        dot_replacements = [(row[0],row[1]) for row in reader]
                    
                    
                    for dot_original, dot_replacement in dot_replacements:
                        r_dot_original=re.escape(dot_original)
                        dot_new_original = rf"\b{r_dot_original}"
                        text = re.sub(dot_new_original, dot_replacement, text)

                segments.extend(split_segment(text.strip(), start_time, end_time))
            segment_index = int(line.strip())
            text = ""
        elif '-->' in line:
            times = line.strip().split(' --> ')
            start_time = convert_time_to_seconds(times[0].replace(',', '.'))
            end_time = convert_time_to_seconds(times[1].replace(',', '.'))
        else:
            text += line.strip() + " "

    if text:
        #text = replace_special_periods(text)
        if st.session_state.select_dot_result:
            
            data_folder = "/content/my_app"
            with open(os.path.join(data_folder,st.session_state.select_dot_result), newline='',encoding='utf-8') as dot_csvfile:
                reader = csv.reader(dot_csvfile)
                next(reader)
                dot_replacements = [(row[0],row[1]) for row in reader]
            
            
            for dot_original, dot_replacement in dot_replacements:
                r_dot_original=re.escape(dot_original)
                dot_new_original = rf"\b{r_dot_original}"
                text = re.sub(dot_new_original, dot_replacement, text)

        segments.extend(split_segment(text.strip(), start_time, end_time))

    merged_segments = merge_segments(segments)

    # replacements.csv を読み込む
    
    if replace_word==True:
        common_app_data = os.getenv('APPDATA')  # WindowsのCommon AppDataフォルダ
        data_folder = "/content/my_app"
        with open(os.path.join(data_folder,st.session_state.select_rp_result), newline='', encoding='utf-8') as csvfile:
            reader = csv.reader(csvfile)
            next(reader)  # ヘッダーをスキップ
            replacements = [(row[0], row[1]) for row in reader]

    output = []
    segment_number = 0
    for text, start, end in merged_segments:
        text = restore_special_periods(text)
        if replace_word==True:
            for original, replacement in replacements:
                original=re.escape(original)
                new_original = rf"\b{original}\b"
                text = re.sub(new_original, replacement, text)
        output.append(f"{segment_number + 1}\n{convert_seconds_to_time(start,'srt').replace('.', ',')} --> {convert_seconds_to_time(end,'srt').replace('.', ',')}\n{text}\n")
        segment_number += 1

    return '\n'.join(output)

def convert_time_to_seconds(time_str):
    try:
        time_parts = time_str.split(':')
        if len(time_parts) == 3:
            h, m, s = time_parts
        elif len(time_parts) == 2:
            h = 0
            m, s = time_parts
        else:
            raise ValueError(f"Unexpected time format: {time_str}")
        
        h = float(h)
        m = float(m)
        s = float(s.replace(',', '.'))
        return h * 3600 + m * 60 + s
    except ValueError as e:
        raise ValueError(f"Error converting time: {time_str}, {e}")

def convert_seconds_to_time(seconds, format_type='vtt'):
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    s = seconds % 60
    if format_type == 'vtt':
        return f"{h:01}:{m:02}:{s:06.3f}".replace(',', '.')
    else:
        return f"{h:02}:{m:02}:{s:06.3f}".replace('.', ',')

def process_file(input_file,replace_word):
    if input_file is None:
        return None, None, [],pd.DataFrame({'1': [''], '2': [''], '3': ['']})#0729



    with open(input_file, 'r', encoding='utf-8') as file:
        lines = file.readlines()

    basename = os.path.splitext(os.path.basename(input_file))[0]
    
    timestamp_patch = datetime.now().strftime("%Y%m%d%H%M%S")
    temp_dir = os.path.join(tempfile.gettempdir(), f"tempdir_{timestamp_patch}")
    os.makedirs(temp_dir, exist_ok=True) 
    if input_file.lower().endswith('.vtt'):
        output = process_vtt(lines,replace_word)
        output_file = os.path.join(temp_dir, f"{basename}_ed.vtt")
    elif input_file.lower().endswith('.srt'):
        output = process_srt(lines,replace_word)
        output_file = os.path.join(temp_dir, f"{basename}_ed.srt")
    else:
        raise ValueError('Unsupported file format')

    with open(output_file, 'w', encoding='utf-8') as file:
        file.write(output)

    with open(output_file,"r",encoding="utf-8") as f:
        for_text=f.read()



    if input_file.endswith(".srt"):
        segment_pattern = r'(\d+)\n(\d{2}:\d{2}:\d{2},\d{3}\s*-->\s*\d{2}:\d{2}:\d{2},\d{3})\n(.*?)(?=\n\d+\n|\Z)'
        segments = re.findall(segment_pattern, for_text , re.DOTALL)  # 全てのセグメントを抽出
                    
        result = []

        # 各セグメントごとにID、タイムスタンプ、テキストを処理
        for segment in segments:
            segment_id, timestamp, text = segment
            
            text=text.replace("\t","").replace("\n"," ")
            result.append(text)

        final_texts=''.join(result) 
    
    elif input_file.endswith(".vtt"):

        t7.webvtt_remover(for_text)
        segment_pattern = r'(\d+)\n(\d{1,2}:\d{2}:\d{2}.\d{3}\s*-->\s*\d{1,2}:\d{2}:\d{2}.\d{3})\n(.*?)(?=\n\d+\n|\Z)'
        segments = re.findall(segment_pattern, for_text, re.DOTALL)  # 全てのセグメントを抽出
        
        result = []

        # 各セグメントごとにID、タイムスタンプ、テキストを処理
        for segment in segments:
            segment_id, timestamp, text = segment
            text=text.replace("\t","").replace("\n"," ")
            result.append(text)   
        final_texts=''.join(result) 
    
  

    timestamp_patch = datetime.now().strftime("%Y%m%d%H%M%S")
    temp_dir = os.path.join(tempfile.gettempdir(), f"tempdir_{timestamp_patch}")
    os.makedirs(temp_dir, exist_ok=True) 

    text_file_path=os.path.join(temp_dir,f"{basename}_ed_NR.txt")
    with open(text_file_path,"w",encoding="utf-8") as f:
        f.write(final_texts)



    t7_excel_file,t7_df=create_excel_from_srt(english_path=output_file,tail="")#0729
    output_html = f"""<head><meta charset="UTF-8"></head><body><pre style="white-space: pre-wrap; overflow-y: auto; height: 300px; word-wrap: break-word; padding: 10px; font-family: inherit; font-size: inherit;">{output}</pre></body>"""


    # return output_html, [output_file, docx_file], output_file
    return output_html, [output_file,t7_excel_file,text_file_path],output_file,t7_df

def correct_srt_format_from_text(text):
    if st.session_state.language_result in ["Japanese (ja)","Chinese (zh)"]:
        print("true")
        text=co.kanji_henkan(text)
    content = re.sub(r'[\u200B-\u200D\uFEFF]', '', text)
    content = content.replace('\u200B',"")
    content = content.strip()
    content = content.replace(" ","^").replace("^^","^")
    content = content.replace("\n","")
    content = re.sub(r'\s+', '', content)
    patterns_replacements = [
        (r'(\d)\^+(\d)', r'\1\2'),  # ① 数字と数字の間
        (r'(\d)\^+(,)', r'\1\2'),   # ④ 数字とカンマの間
        (r'(,)\^+(\d)', r'\1\2'),   # ⑤ カンマと数字の間
        (r'(\d)\^+(:)', r'\1\2'),   # ⑥ 数字とコロンの間
        (r'(:)\^+(\d)', r'\1\2'),   # ⑦ コロンと数字の間
    ]    
    for pattern, replacement in patterns_replacements:
        content = re.sub(pattern, replacement, content)    



    pattern = re.compile(r'(\d{1,4})\^*(\d{2}:\d{2}:\d{2},\d{3})\^*-->\^*(\d{2}:\d{2}:\d{2},\d{3})')
    matches = pattern.findall(content)
    segments = pattern.split(content)
    #print(f"segments:{segments}")
    corrected_content = []
    for i in range(1, len(segments), 4):
        segment_id = segments[i].replace("^","")
        start_timestamp = segments[i + 1].replace("^","")
        end_timestamp = segments[i + 2].replace("^","")
        segment_text = segments[i + 3].replace("^"," ").strip()

        corrected_content.append(f"{segment_id}")
        corrected_content.append(f"{start_timestamp} --> {end_timestamp}")
        corrected_content.append(segment_text)

    final_content = "\n\n".join("\n".join(block) for block in zip(*[iter(corrected_content)]*3))
    return final_content

def correct_vtt_format_from_text(text):
    if st.session_state.language_result in ["Japanese (ja)","Chinese (zh)"]:
        text=co.kanji_henkan(text)
    content = re.sub(r'[\u200B-\u200D\uFEFF]', '', text)
    content = content.replace('\u200B',"")
    content = content.replace(" ","^").replace("^^","^")
    content = content.replace("\n","")
    content = re.sub(r'\s+','',content)
    #print(content)
    patterns_replacements = [
        (r'(\d)\^+(\d)', r'\1\2'),  # ① 数字と数字の間
        (r'(\d)\^+(\.)', r'\1\2'),  # ② 数字とピリオドの間
        (r'(\.)\^+(\d)', r'\1\2'),  # ③ ピリオドと数字の間
        (r'(\d)\^+(:)', r'\1\2'),   # ⑥ 数字とコロンの間
        (r'(:)\^+(\d)', r'\1\2'),   # ⑦ コロンと数字の間
    ]    
    #print(content)
    for pattern, replacement in patterns_replacements:
        content = re.sub(pattern, replacement, content)    
    
    pattern = re.compile(r'(\d{1,4})\^*(\d{1}:\d{2}:\d{2}\.\d{3})\^*-->\^*(\d{1}:\d{2}:\d{2}\.\d{3})')
    matches = pattern.findall(content)
    if not matches:
        pattern = re.compile(r'(\d{1,4})\^*(\d{2}:\d{2}:\d{2}\.\d{3})\^*-->\^*(\d{2}:\d{2}:\d{2}\.\d{3})')
        matches = pattern.findall(content)

    segments = pattern.split(content)
    #print(segments)
    corrected_content = []
    for i in range(1, len(segments), 4):
        segment_id = segments[i].replace("^","")
        start_time = segments[i + 1].replace("^","")
        end_time = segments[i + 2].replace("^","")
        text = segments[i + 3].replace("^"," ").strip()
        corrected_content.append(f"{segment_id}")
        corrected_content.append(f"{start_time} --> {end_time}")
        corrected_content.append(text)
    final_content = "WEBVTT\n\n" + "\n\n".join("\n".join(block) for block in zip(*[iter(corrected_content)]*3))
    return final_content

def vtt_translate(input_file, translated_content, output_file):
    if input_file is None or translated_content is None or output_file is None:
        return None, None
    basename = os.path.splitext(os.path.basename(input_file))[0]
  

    timestamp_patch = datetime.now().strftime("%Y%m%d%H%M%S")
    temp_dir = os.path.join(tempfile.gettempdir(), f"tempdir_{timestamp_patch}")
    os.makedirs(temp_dir, exist_ok=True) 
    lang_tail=co.extract_short_name(st.session_state.language_result)
    ja_file_name, file_extension = os.path.splitext(input_file)
    output_ja_file_path = os.path.join(temp_dir, f"{basename}_ed_{lang_tail}{file_extension}")
    if file_extension == ".srt":
        corrected_content = correct_srt_format_from_text(translated_content)
    elif file_extension == ".vtt":
        corrected_content = correct_vtt_format_from_text(translated_content)
        
    with open(output_ja_file_path, 'w', encoding='utf-8') as file:
        file.write(corrected_content + '\n')
    output_excel_file, df_data = create_excel(output_file, output_ja_file_path)
    if st.session_state.ja_split:
        
        lang_tail=co.extract_short_name(st.session_state.language_result)
        print(f"lang_tail:{lang_tail}")
        if lang_tail in ["ja","zh"]:
            print("ja_split_true")
            splitted_file=t8.process_files([output_ja_file_path],True)
        else:
            splitted_file=period.process_files([output_ja_file_path],True)

        
        if st.session_state.comma_split:
            #print("commmanow")
            print(st.session_state.max_split)
            splitted_cm_path=cm.split_srt_vtt_by_comma_and_merge(splitted_file,st.session_state.max_split,st.session_state.min_split)
            return [output_ja_file_path, splitted_file,splitted_cm_path,output_excel_file], df_data                          
                                    
        return [output_ja_file_path, splitted_file,output_excel_file], df_data
    return [output_ja_file_path, output_excel_file], df_data

def webvtt_remover(sentence):
    sentence = re.sub(r'[\u200B-\u200D\uFEFF]', '', sentence)
    sentence = sentence.replace('\r\n', '\n').replace('\r', '\n')
    sentence = sentence.replace('\u200B',"")

    pattern1 = re.compile(r'(\d+)\n(\d{1}:\d{2}:\d{2}\.\d{3}-->\d{1}:\d{2}:\d{2}\.\d{3})', re.DOTALL)
    pattern2 = re.compile(r'(\d+)\n(\d{2}:\d{2}:\d{2}\.\d{3}-->\d{2}:\d{2}:\d{2}\.\d{3})', re.DOTALL)

    match = pattern1.search(sentence)
    if match:
        start_index = match.start(1)
        rm_webvtt_sentence = sentence[start_index:]
    else:
        match = pattern2.search(sentence)
        if match:
            start_index = match.start(1)
            rm_webvtt_sentence = sentence[start_index:]
        else:
            rm_webvtt_sentence = sentence

    return rm_webvtt_sentence

def webvtt_remover_NR(sentence):
    sentence = re.sub(r'[\u200B-\u200D\uFEFF]', '', sentence)
    sentence = sentence.replace('\r\n', '\n').replace('\r', '\n')
    sentence = sentence.replace('\u200B',"")

    pattern1 = re.compile(r'(\d+)(\d{1}:\d{2}:\d{2}\.\d{3}-->\d{1}:\d{2}:\d{2}\.\d{3})', re.DOTALL)
    pattern2 = re.compile(r'(\d+)(\d{2}:\d{2}:\d{2}\.\d{3}-->\d{2}:\d{2}:\d{2}\.\d{3})', re.DOTALL)

    match = pattern1.search(sentence)
    if match:
        start_index = match.start(1)
        rm_webvtt_sentence = sentence[start_index:]
    else:
        match = pattern2.search(sentence)
        if match:
            start_index = match.start(1)
            rm_webvtt_sentence = sentence[start_index:]
        else:
            rm_webvtt_sentence = sentence

    return rm_webvtt_sentence

def create_excel(output_file, output_ja_file_path,REP=False):
    segments = []
    if output_file.lower().endswith('.vtt') or output_file.lower().endswith('.srt'):
        with open(output_file, 'r',encoding="utf-8") as f:
            lines = f.readlines()
        if output_file.lower().endswith('.vtt'):
            lines = webvtt_remover(''.join(lines))
            #print(st.session_state.replace_word,REP)
            if st.session_state.replace_word and REP==True:

                common_app_data = os.getenv('APPDATA')  # WindowsのCommon AppDataフォルダ
                data_folder = "/content/my_app"
                with open(os.path.join(data_folder,st.session_state.select_rp_result), newline='', encoding='utf-8') as csvfile:
                    reader = csv.reader(csvfile)
                    next(reader)  # ヘッダーをスキップ
                    replacements = [(row[0], row[1]) for row in reader]

                for original, replacement in replacements:
                    original=re.escape(original)
                    new_original = rf"\b{original}\b"   
                    lines=re.sub(new_original, replacement, lines)
            lines=lines.splitlines()
        else:                         
            #print(st.session_state.replace_word,REP)
            lines=''.join(lines)
            if st.session_state.replace_word and REP==True:
                common_app_data = os.getenv('APPDATA')  # WindowsのCommon AppDataフォルダ
                data_folder = "/content/my_app"
                with open(os.path.join(data_folder,st.session_state.select_rp_result), newline='', encoding='utf-8') as csvfile:
                    reader = csv.reader(csvfile)
                    next(reader)  # ヘッダーをスキップ
                    replacements = [(row[0], row[1]) for row in reader]

                for original, replacement in replacements:
                    original=re.escape(original)
                    new_original = rf"\b{original}\b"   
                    lines=re.sub(new_original, replacement, lines)

            lines=lines.splitlines() 

        segments.extend(parse_segments(lines))


    if output_ja_file_path.lower().endswith('.vtt') or output_ja_file_path.lower().endswith('.srt'):
        with codecs.open(output_ja_file_path, 'r', 'utf-8') as f:
            ja_lines = f.readlines()
            if output_ja_file_path.lower().endswith('.vtt'):
                ja_lines = webvtt_remover(''.join(ja_lines)).splitlines()
            ja_segments = parse_segments(ja_lines)




    excel_data = []
    lang_tail=co.extract_short_name(st.session_state.language_result)
    lang_name=(st.session_state.language_result).replace(f" ({lang_tail})","")
    if output_ja_file_path.lower().endswith('.vtt'):
        for (eng_segment, start, end), ja_segment in zip(segments, ja_segments):
            if start is not None and end is not None:
                excel_data.append({
                    'ID': segments.index((eng_segment, start, end)) + 1,
                    'Start': convert_seconds_to_time(start, 'vtt'),
                    'End': convert_seconds_to_time(end, 'vtt'),
                    'English Subtitle': eng_segment,
                    f'{lang_name} Subtitle': ja_segment[0]
                })
            else:
                print(f"Skipping segment due to missing time: {eng_segment}")
    else:
        for (eng_segment, start, end), ja_segment in zip(segments, ja_segments):
            if start is not None and end is not None:
                excel_data.append({
                    'ID': segments.index((eng_segment, start, end)) + 1,
                    'Start': convert_seconds_to_time(start, 'srt'),
                    'End': convert_seconds_to_time(end, 'srt'),
                    'English Subtitle': eng_segment,
                    f'{lang_name} Subtitle': ja_segment[0]
                })
            else:
                print(f"Skipping segment due to missing time: {eng_segment}")

    df_data = pd.DataFrame(excel_data)


    timestamp_patch = datetime.now().strftime("%Y%m%d%H%M%S")
    temp_dir = os.path.join(tempfile.gettempdir(), f"tempdir_{timestamp_patch}")
    os.makedirs(temp_dir, exist_ok=True) 

    output_excel_file = os.path.join(temp_dir, f"{os.path.splitext(os.path.basename(output_file))[0]}.xlsx")
    df_data.to_excel(output_excel_file, index=False)
    with pd.ExcelWriter(output_excel_file, engine='openpyxl') as writer:
        df_data.to_excel(writer, index=False, sheet_name='Subtitles')
        workbook = writer.book
        worksheet = writer.sheets['Subtitles']
        column_widths = {'A': 7, 'B': 25, 'C': 25, 'D': 90, 'E': 90}
        for column, width in column_widths.items():
            worksheet.column_dimensions[column].width = width
        for row in worksheet.iter_rows(min_row=2, max_row=len(df_data) + 1):
            for cell in row:
                if cell.column_letter == 'A':
                    cell.alignment = Alignment(horizontal='right', vertical='center')
                elif cell.column_letter in ['B', 'C']:
                    cell.alignment = Alignment(horizontal='center', vertical='center')
                elif cell.column_letter in ['D', 'E']:
                    cell.alignment = Alignment(horizontal='left', vertical='center')
        for row in worksheet.iter_rows(min_row=2, max_row=len(df_data) + 1):
            worksheet.row_dimensions[row[0].row].height = 30
        header_font = Font(bold=True)
        for cell in worksheet["1:1"]:
            cell.font = header_font
            cell.alignment = Alignment(horizontal='center')
            cell.fill = PatternFill(start_color="DAEEF3", end_color="DAEEF3", fill_type="solid")
    return output_excel_file, df_data

def parse_segments(lines):
    segments = []
    segment_text = ""
    start_time = None
    end_time = None

    timestamp_pattern = re.compile(r'(\d{1,2}:\d{2}:\d{2}[.,]\d{3}) --> (\d{1,2}:\d{2}:\d{2}[.,]\d{3})')

    for line in lines:
        line = line.strip()
        if timestamp_pattern.search(line):
            if segment_text and start_time and end_time:
                segments.append((segment_text.strip(), start_time, end_time))
            segment_text = ""
            match = timestamp_pattern.search(line)
            start_time = convert_time_to_seconds(match.group(1).replace(',', '.'))
            end_time = convert_time_to_seconds(match.group(2).replace(',', '.'))
        elif re.match(r'^\d+$', line):
            continue  # Ignore segment numbers
        else:
            segment_text += line + " "

    if segment_text and start_time and end_time:
        segments.append((segment_text.strip(), start_time, end_time))

    return segments



## tab2
def read_file_content_from_path(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        content = file.read()
    return content

def read_file_content(file):
    if file is None:
        return """<div style='color: orange !important; font-family: inherit; text-align: center; 
                display: flex; align-items: flex-start; justify-content: center; height: 400px; padding-top: 40px;'>
                No file uploaded
                </div>"""

    file_extension = os.path.splitext(file)[1]
    content = ""

   
    if file_extension == '.txt':
        content = read_file_content_from_path(file)
        content = f"""<pre style="white-space: pre-wrap; overflow-y: auto; height: 300px; word-wrap: break-word; padding: 10px; font-family: inherit; font-size: inherit;">{content}</pre>"""

    elif file_extension == '.srt':
        content = read_file_content_from_path(file)
        content = unify_timestamps(content)
        content = f"""<pre style="white-space: pre-wrap; overflow-y: auto; height: 300px; word-wrap: break-word; padding: 10px; font-family: inherit; font-size: inherit;">{content}</pre>"""

    elif file_extension == '.vtt':
        content = read_file_content_from_path(file)
        content = unify_timestamps_vtt(content)
        content = f"""<pre style="white-space: pre-wrap; overflow-y: auto; height: 300px; word-wrap: break-word; padding: 10px; font-family: inherit; font-size: inherit;">{content}</pre>"""
    else:
        return None
    return content

def display_file_content(file,replace_word=False):
    if file is None:
        return read_file_content(file),[]
    
    file_content_for_html=read_file_content(file)
    #print(file_content_for_html)
    if replace_word:
        common_app_data = os.getenv('APPDATA')  # WindowsのCommon AppDataフォルダ
        data_folder = "/content/my_app"
        with open(os.path.join(data_folder,st.session_state.select_rp_result), newline='', encoding='utf-8') as csvfile:
            reader = csv.reader(csvfile)
            next(reader)  # ヘッダーをスキップ
            replacements = [(row[0], row[1]) for row in reader]
            #print(replacements)

        for original, replacement in replacements:
            original=re.escape(original)
            new_original = rf"\b{original}\b"   
            file_content_for_html=re.sub(new_original, replacement, file_content_for_html)
            
    if file.endswith('.txt'):
        filename = os.path.basename(file)    
        match = re.match(r"(.+?)(_NR\.txt|_R\.txt)$", filename)
        
        if match:
            basename, ext = match.groups()
            timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
            temp_dir = os.path.join(tempfile.gettempdir(), f"tempdir_{timestamp}")
            os.makedirs(temp_dir, exist_ok=True)
            
            if ext == "_NR.txt":
                doc_filename = os.path.join(temp_dir, f"{basename}_txtnr.docx")
            elif ext == "_R.txt":
                doc_filename = os.path.join(temp_dir, f"{basename}_txtr.docx")
      
            with open(file, 'r', encoding='utf-8') as f:
                content = f.read()
            

            if replace_word:
                for original, replacement in replacements:
                    original=re.escape(original)
                    new_original = rf"\b{original}\b"
                    content = re.sub(new_original, replacement, content)
            doc = Document()
            doc.add_paragraph(content)
            doc.save(doc_filename)

            return file_content_for_html, doc_filename #pd.DataFrame({'1': [''], '2': [''], '3': ['']})
        
        else:
            #print("txtです。")
            st.error('テキストファイルは"_NR.txt"または"_R.txt"の形式のファイルをアップロードしてください。')
            return "",None
              
    else:
        if file.endswith('.srt'):
            excel_path,t4_df = t3.create_excel_from_srt(english_path=file,tail="",replace_word=replace_word)
        elif file.endswith('.vtt'):
            excel_path,t4_df = t3.create_excel_from_srt(english_path=file,tail="",replace_word=replace_word)


        return file_content_for_html, excel_path #t4_df

def save_translated_content(file, translated_text):
    if file is None:
        return []
    
    file_name, file_extension = os.path.splitext(os.path.basename(file))

    timestamp_patch = datetime.now().strftime("%Y%m%d%H%M%S")
    temp_dir = os.path.join(tempfile.gettempdir(), f"tempdir_{timestamp_patch}")
    os.makedirs(temp_dir, exist_ok=True) 
    lang_tail=co.extract_short_name(st.session_state.language_result)
    output_file_path = os.path.join(temp_dir, f"{file_name}_{lang_tail}{file_extension}")

    '''if file_extension == '.docx':
        doc = Document()
        doc.add_paragraph(translated_text)
        doc.save(output_file_path)'''

    if file_extension == '.txt':
        with codecs.open(output_file_path, 'w', 'utf-8') as f:
            f.write(translated_text)
        
        return[output_file_path],None
    
    elif file_extension in ['.srt', '.vtt']:
        if st.session_state.language_result in ["Japanese (ja)","Chinese (zh)"]:
            translated_text=co.kanji_henkan(translated_text)
        content = re.sub(r'[\u200B-\u200D\uFEFF]', '', translated_text)
        content = content.replace('\u200B',"")
        content = content.replace(" ","^").replace("^^","^")
        content = content.replace("\n","")
        content = re.sub(r'\s+', '', content)

        

        if file_extension == '.srt':
            patterns_replacements = [
                (r'(\d)\^+(\d)', r'\1\2'),  # ① 数字と数字の間
                (r'(\d)\^+(,)', r'\1\2'),   # ④ 数字とカンマの間
                (r'(,)\^+(\d)', r'\1\2'),   # ⑤ カンマと数字の間
                (r'(\d)\^+(:)', r'\1\2'),   # ⑥ 数字とコロンの間
                (r'(:)\^+(\d)', r'\1\2'),   # ⑦ コロンと数字の間
            ]
            for pattern, replacement in patterns_replacements:
                content = re.sub(pattern, replacement, content)
            

            pattern = re.compile(r'(\d{1,4})\^*(\d{2}:\d{2}:\d{2},\d{3}\^*-->\^*\d{2}:\d{2}:\d{2},\d{3})')

        elif file_extension == '.vtt':
            patterns_replacements = [
                (r'(\d)\^+(\d)', r'\1\2'),  # ① 数字と数字の間
                (r'(\d)\^+(\.)', r'\1\2'),  # ② 数字とピリオドの間
                (r'(\.)\^+(\d)', r'\1\2'),  # ③ ピリオドと数字の間
                (r'(\d)\^+(:)', r'\1\2'),   # ⑥ 数字とコロンの間
                (r'(:)\^+(\d)', r'\1\2'),   # ⑦ コロンと数字の間
            ]
            for pattern, replacement in patterns_replacements:
                content = re.sub(pattern, replacement, content)

            
            pattern = re.compile(r'(\d{1,4})\^*(\d{1}:\d{2}:\d{2}\.\d{3}\^*-->\^*\d{1}:\d{2}:\d{2}\.\d{3})')
        
        matches = pattern.findall(content)        
        if not matches:
            pattern =  re.compile(r'(\d{1,4})\^*(\d{2}:\d{2}:\d{2}\.\d{3}\^*-->\^*\d{2}:\d{2}:\d{2}\.\d{3})')
            matches = pattern.findall(content)

        if not matches:
            pattern =  re.compile(r'(\d{1,4})\^*(\d{1}:\d{2}:\d{2}\,\d{3}\^*-->\^*\d{1}:\d{2}:\d{2}\,\d{3})')
            matches = pattern.findall(content)

        segments = pattern.split(content)
        corrected_content = []
        
        for i in range(1, len(segments), 3):
            segment_id = segments[i].replace("^","").strip()
            timestamp = segments[i + 1].replace("^","").strip()
            text = segments[i + 2].replace("^"," ").strip()
            
            corrected_content.append(f"{segment_id}")
            corrected_content.append(timestamp.replace('-->', ' --> '))
            corrected_content.append(text)
        if file_extension==".vtt":
            corrected_content ="WEBVTT\n\n"+"\n\n".join("\n".join(block) for block in zip(*[iter(corrected_content)]*3))        
        else:
            corrected_content ="\n\n".join("\n".join(block) for block in zip(*[iter(corrected_content)]*3))        
        
        with open(output_file_path, 'w', encoding='utf-8') as f:
            f.write(corrected_content + '\n')

    output_excel_file, df_data = create_excel(file, output_file_path,True)
    if st.session_state.ja_split==True:
        lang_tail=co.extract_short_name(st.session_state.language_result)
        if lang_tail=='ja' or lang_tail=='zh':
            splitted_file=t8.process_files([output_file_path],True)
        else:
            splitted_file=period.process_files([output_file_path],True)
            print("not ja group")
        if st.session_state.comma_split:
            splitted_cm_file=cm.split_srt_vtt_by_comma_and_merge(splitted_file,st.session_state.max_split,st.session_state.min_split)
            return [output_file_path,splitted_file,splitted_cm_file,output_excel_file],df_data
        return [output_file_path,splitted_file,output_excel_file],df_data
        

    #print(df_data)
    return [output_file_path,output_excel_file],df_data

#tab3

def unify_timestamps_vtt(text):
    # Define patterns for different timestamp formats
    pattern_1_digit = re.compile(r'(\d{2}:\d{2}:\d{2}\.\d)(?!\d)')
    pattern_2_digits = re.compile(r'(\d{2}:\d{2}:\d{2}\.\d{2})(?!\d)')
    pattern_3_digit = re.compile(r'(\d{1}:\d{2}:\d{2}\.\d)(?!\d)')
    pattern_4_digits = re.compile(r'(\d{1}:\d{2}:\d{2}\.\d{2})(?!\d)')

    # Replace 1-digit and 2-digit millisecond formats with 3-digit format for HH:MM:SS format
    if pattern_1_digit.search(text):
        text = pattern_1_digit.sub(lambda x: x.group(1) + '00', text)
        #print("pattern1")
    if pattern_2_digits.search(text):
        text = pattern_2_digits.sub(lambda x: x.group(1) + '0', text)
        #print("pattern2")

    # Replace 1-digit and 2-digit millisecond formats with 3-digit format for H:MM:SS format
    if pattern_3_digit.search(text):
        text = pattern_3_digit.sub(lambda x: x.group(1) + '00', text)
        #print("pattern3")
    if pattern_4_digits.search(text):
        text = pattern_4_digits.sub(lambda x: x.group(1) + '0', text)
        #print("pattern4")

    return text


#srtファイルのタイムスタンプ桁数を統一。
def unify_timestamps(text):
    # Define patterns for different timestamp formats
    pattern_1_digit = re.compile(r'(\d{2}:\d{2}:\d{2}\,\d)(?!\d)')
    pattern_2_digits = re.compile(r'(\d{2}:\d{2}:\d{2}\,\d{2})(?!\d)')
    
    # Replace 1-digit millisecond format with 3-digit format
    text = pattern_1_digit.sub(lambda x: x.group(1) + '00', text)
    
    # Replace 2-digit millisecond format with 3-digit format
    text = pattern_2_digits.sub(lambda x: x.group(1) + '0', text)
    
    return text

def unify_timestamps_forlist(lines, format_type):
    if format_type == 'vtt':
        unify_timestamps_func = unify_timestamps_vtt
    elif format_type == 'srt':
        unify_timestamps_func = unify_timestamps
    else:
        raise ValueError(f"Unsupported format type: {format_type}")

    return [unify_timestamps_func(line) if '-->' in line else line for line in lines]

def convert_docx_to_srttxt(docx_files):
    output_files = []
    if docx_files is None:
        return []
    for docx_file in docx_files:
        try:
           
            filename = os.path.basename(docx_file)
            base_name, ext = os.path.splitext(filename)
            clean_name = re.sub(r'[\r\n]+', '', base_name)
            #print(f"Processing file: {clean_name}")
            lang_tail=co.extract_short_name(st.session_state.language_result)
            if clean_name.endswith('_srt(1)'):
                output_filename = clean_name.replace('_srt(1)', f'_{lang_tail}.srt')
            elif clean_name.endswith("_srt"):
                output_filename = clean_name.replace("_srt", f"_{lang_tail}.srt")
            elif clean_name.endswith('_vtt(1)'):
                output_filename = clean_name.replace('_vtt(1)', f'_{lang_tail}.vtt')
            elif clean_name.endswith("_vtt"):
                output_filename = clean_name.replace("_vtt", f"_{lang_tail}.vtt")
            elif clean_name.endswith('_txtnr(1)'):
                output_filename = clean_name.replace('_txtnr(1)', f'_NR_{lang_tail}.txt')
            elif clean_name.endswith("_txtnr"):
                output_filename = clean_name.replace("_txtnr", f"_NR_{lang_tail}.txt")
            elif clean_name.endswith('_txtr(1)'):
                output_filename = clean_name.replace('_txtr(1)', f'_R_{lang_tail}.txt')
            elif clean_name.endswith("_txtr"):
                output_filename = clean_name.replace("_txtr", f"_R_{lang_tail}.txt")
            else:
                print(f"Skipping file with unrecognized pattern: {clean_name}")
                continue
            timestamp_patch = datetime.now().strftime("%Y%m%d%H%M%S")
            temp_dir = os.path.join(tempfile.gettempdir(), f"tempdir_{timestamp_patch}")
            os.makedirs(temp_dir, exist_ok=True) 
            output_filepath = os.path.join(temp_dir, f"{output_filename}")
            #if clean_name.endswith('_NR_ja.txt') or clean_name.endswith('_R_ja.txt'):

 
            
            if output_filename.endswith('.txt'):
                doc = Document(docx_file)
                content = "\n".join([para.text for para in doc.paragraphs])
                #print(f"Initial content read from file: {content[:100]}")  # Show only the first 100 characters for brevity

            if output_filename.endswith('.srt'):
                output_filepath=t5.excel_to_srt(docx_file,output_filepath)

            elif output_filename.endswith('vtt'):
                output_filepath=t5.excel_to_vtt(docx_file,output_filepath)

            
            if output_filename.endswith('.txt'):
                final_content = content

                with open(output_filepath, 'w', encoding='utf-8') as output_file:
                    output_file.write(final_content)
            
            output_files.append(output_filepath)
        except Exception as e:
            print(f"An error occurred while processing {filename}: {str(e)}")

    if len(output_files) > 3:
        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
        temp_dir = os.path.join(tempfile.gettempdir(), f"tempdir_{timestamp}")
        os.makedirs(temp_dir, exist_ok=True)
        zip_filename = os.path.join(temp_dir, f"converted_from_docx_{lang_tail}.zip")
        with zipfile.ZipFile(zip_filename, 'w') as zipf:
            for file in output_files:
                zipf.write(file, os.path.basename(file))
        
        return [zip_filename]
    
    return output_files



def process_doc_files(files,replace_word=False):
    output_files = []
    if not files:
        return []

    for file in files:

        filename = os.path.basename(file)
        match = re.match(r"(.+?)(_NR\.txt|_R\.txt|\.srt|\.vtt)$", filename)
        if not match:
            st.error('ファイル末尾が、"_NR.txt", "_R.txt"に該当しないtextファイルはスキップされました。')
            continue  # skip files with unknown extensions
        
        basename, ext = match.groups()
        if ext == ".srt":
            doc_filename = f"{basename}_srt.xlsx"
        elif ext == ".vtt":
            doc_filename = f"{basename}_vtt.xlsx"
        elif ext == "_NR.txt":
            doc_filename = f"{basename}_txtnr.docx"
        elif ext == "_R.txt":
            doc_filename = f"{basename}_txtr.docx"

        if ext in ['.srt', '.vtt']:
            with open(file, 'r', encoding='utf-8') as f:
                content = f.read()
            if replace_word==True:
                common_app_data = os.getenv('APPDATA')  # WindowsのCommon AppDataフォルダ
                data_folder = "/content/my_app"
                with open(os.path.join(data_folder,st.session_state.select_rp_result), newline='', encoding='utf-8') as csvfile:
                    reader = csv.reader(csvfile)
                    next(reader)  # ヘッダーをスキップ
                    replacements = [(row[0], row[1]) for row in reader]
                for original, replacement in replacements:
                    original=re.escape(original)
                    new_original = rf"\b{original}\b" 
                    content = re.sub(new_original, replacement, content)
            #print(content)
            
            if ext == '.srt':
                unified_content = t4.unify_timestamps(content)
                parsed_content = parse_srt_cc(unified_content)
            else:
                unified_content = t4.unify_timestamps_vtt(content)
                parsed_content = parse_vtt_cc(unified_content)

            t5_excel_path, t5_df = create_excel_from_srt_cc(parsed_content, doc_filename)
            output_files.append(t5_excel_path)
            #print(f"Processed file: {t5_excel_path}")
        
        elif ext in ["_NR.txt", "_R.txt"]:
            with open(file, 'r', encoding='utf-8') as f:
                content = f.read()
            doc = Document()
            doc.add_paragraph(content)
            timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
            temp_dir = os.path.join(tempfile.gettempdir(), f"tempdir_{timestamp}")
            os.makedirs(temp_dir, exist_ok=True)
            doc_path=os.path.join(temp_dir,doc_filename)
            doc.save(doc_path)
            output_files.append(doc_path)
    
    if len(output_files) > 3:
        zip_filename = "converted_from_srttxt_en.zip"
        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
        temp_dir = os.path.join(tempfile.gettempdir(), f"tempdir_{timestamp}")
        os.makedirs(temp_dir, exist_ok=True)
        new_zip_path=os.path.join(temp_dir,zip_filename)
        with zipfile.ZipFile(new_zip_path, 'w') as zip_file:
            for file in output_files:
                zip_file.write(file, os.path.basename(file))
        
        return [new_zip_path]
    
    return output_files

def parse_srt_cc(srt_content):
    pattern = r'(\d+)\n(\d{2}:\d{2}:\d{2},\d{3}) --> (\d{2}:\d{2}:\d{2},\d{3})\n(.*?)(?:\n\n|\Z)'
    matches = re.findall(pattern, srt_content, re.DOTALL)
    
    subtitles = []
    for match in matches:
        subtitles.append({
            'ID': int(match[0]),
            'Start': match[1],
            'End': match[2],
            'Text': match[3].replace('\n', ' ')
        })
    
    return subtitles

def parse_vtt_cc(vtt_content):
    pattern = r'(\d+)\n(\d{2}:\d{2}:\d{2}\.\d{3}) --> (\d{2}:\d{2}:\d{2}\.\d{3})\n(.*?)(?:\n\n|\Z)'
    matches = re.findall(pattern, vtt_content, re.DOTALL)
    if not matches:
        pattern = r'(\d+)\n(\d{1}:\d{2}:\d{2}\.\d{3}) --> (\d{1}:\d{2}:\d{2}\.\d{3})\n(.*?)(?:\n\n|\Z)'
        matches = re.findall(pattern, vtt_content, re.DOTALL)
    
    subtitles = []
    for match in matches:
        subtitles.append({
            'ID': int(match[0]),
            'Start': match[1],
            'End': match[2],
            'Text': match[3].replace('\n', ' ')
        })
    
    return subtitles

def create_excel_from_srt_cc(parsed_content, doc_filename):
    df = pd.DataFrame(parsed_content)
    
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    temp_dir = os.path.join(tempfile.gettempdir(), f"tempdir_{timestamp}")
    os.makedirs(temp_dir, exist_ok=True)
    
    excel_path = os.path.join(temp_dir, doc_filename)
    df.to_excel(excel_path, index=False)

    with pd.ExcelWriter(excel_path, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Subtitles')
        workbook = writer.book
        worksheet = writer.sheets['Subtitles']

        column_widths = {'A': 7, 'B': 25, 'C': 25, 'D': 90, 'E': 90}
        for column, width in column_widths.items():
            worksheet.column_dimensions[column].width = width

        for row in worksheet.iter_rows(min_row=2, max_row=len(df) + 1):
            for cell in row:
                if cell.column_letter == 'A':
                    cell.alignment = Alignment(horizontal='right', vertical='center')
                elif cell.column_letter in ['B', 'C']:
                    cell.alignment = Alignment(horizontal='center', vertical='center')
                elif cell.column_letter in ['D', 'E']:
                    cell.alignment = Alignment(horizontal='left', vertical='center')

        for row in worksheet.iter_rows(min_row=2, max_row=len(df) + 1):
            worksheet.row_dimensions[row[0].row].height = 30

        header_font = Font(bold=True)
        for cell in worksheet["1:1"]:
            cell.font = header_font
            cell.alignment = Alignment(horizontal='center')
            cell.fill = PatternFill(start_color="DAEEF3", end_color="DAEEF3", fill_type="solid")

    
    return excel_path, df



    

#tab4
# SRTファイルを解析する関数
def parse_srt(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        lines = file.readlines()
    
    # タイムスタンプの整形を適用
    lines = unify_timestamps_forlist(lines, 'srt')
    
    # 再度文字列として結合
    content = ''.join(lines).replace('\u200B', '')
    pattern = re.compile(r'(\d+)\n(\d{2}:\d{2}:\d{2},\d{3}) --> (\d{2}:\d{2}:\d{2},\d{3})\n(.*?)(?=\n\d|\Z)', re.DOTALL)
    matches = pattern.findall(content)
    #print(f"parse_srt_len(matches):{len(matches)}")

    subtitles = []
    for match in matches:
        subtitles.append({
            'ID': int(match[0]),
            'Start': match[1],
            'End': match[2],
            'Text': match[3].replace('\n', ' ')
        })
    
    return subtitles

def parse_vtt(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        lines = file.readlines()

    # タイムスタンプの整形を適用
    lines = unify_timestamps_forlist(lines, 'vtt')

    # 再度文字列として結合
    content = ''.join(lines).replace('\u200B', '')
    content = webvtt_remover(content)
    #print(f"after_webvtt_remover: {content}")
    
    pattern = re.compile(r'(\d+)\n(\d{1}:\d{2}:\d{2}\.\d{3}) --> (\d{1}:\d{2}:\d{2}\.\d{3})\n(.*?)(?=\n\d|\Z)', re.DOTALL)
    matches = pattern.findall(content)
    #print(f"parse_vtt_d1:Len(matches): {len(matches)}")
    
    if len(matches) == 0:
        pattern = re.compile(r'(\d+)\n(\d{2}:\d{2}:\d{2}\.\d{3}) --> (\d{2}:\d{2}:\d{2}\.\d{3})\n(.*?)(?=\n\d|\Z)', re.DOTALL)
        matches = pattern.findall(content)
        #print(f"parse_vtt_d2:Len(matches): {len(matches)}")

    subtitles = []
    for match in matches:
        subtitles.append({
            'ID': int(match[0]),
            'Start': match[1],
            'End': match[2],
            'Text': match[3].replace('\n', ' ')
        })
    
    return subtitles
def create_excel_from_srt(english_path=None, japanese_path=None,tail="",replace_word=False):
    if replace_word==True:
        common_app_data = os.getenv('APPDATA')  # WindowsのCommon AppDataフォルダ
        data_folder = "/content/my_app"
        with open(os.path.join(data_folder,st.session_state.select_rp_result), newline='', encoding='utf-8') as csvfile:
            reader = csv.reader(csvfile)
            next(reader)  # ヘッダーをスキップ
            replacements = [(row[0], row[1]) for row in reader]
    print(st.session_state.language_result)
    
    lang_tail=co.extract_short_name(st.session_state.language_result)
    lang_name=(st.session_state.language_result).replace(f" ({lang_tail})","")
    if english_path and japanese_path:
        _, file_extension_en = os.path.splitext(english_path)
        _, file_extension_ja = os.path.splitext(japanese_path)

        if file_extension_en.lower() == '.vtt' and file_extension_ja.lower()=='.vtt':
            english_subtitles = parse_vtt(english_path)
            japanese_subtitles = parse_vtt(japanese_path)
        elif file_extension_en=='.srt' and file_extension_ja=='.srt':
            english_subtitles = parse_srt(english_path)
            japanese_subtitles = parse_srt(japanese_path)
        else:
            print('ファイル形式が一致しません')
            return None,None

        data = []
        for eng, jap in zip(english_subtitles, japanese_subtitles):
            if replace_word==True:
                for original, replacement in replacements:
                    original=re.escape(original)
                    new_original = rf"\b{original}\b"
                    eng['Text'] = re.sub(new_original, replacement, eng['Text'])


            data.append({
                'ID': eng['ID'],
                'Start': eng['Start'],
                'End': eng['End'],
                'English Subtitle': eng['Text'],
                f'{lang_name} Subtitle': jap['Text']
            })

        df = pd.DataFrame(data)
        base_name = os.path.splitext(os.path.basename(english_path))[0]
        if file_extension_en=='.srt':
            excel_file_name = f"{base_name}.xlsx"
        elif file_extension_en=='.vtt':
            excel_file_name = f"{base_name}.xlsx"

    elif english_path:
        _, file_extension_en = os.path.splitext(english_path)
        if file_extension_en.lower()=='.srt':
            english_subtitles = parse_srt(english_path)
        elif file_extension_en.lower()=='.vtt':
            english_subtitles = parse_vtt(english_path)
        else:
            print('入力に誤りがあります。')
            return None,None

        data = []
        for eng in english_subtitles:
            if replace_word==True:
                for original, replacement in replacements:
                    original=re.escape(original)
                    new_original = rf"\b{original}\b"
                    eng['Text'] = re.sub(new_original, replacement, eng['Text'])

            data.append({
                'ID': eng['ID'],
                'Start': eng['Start'],
                'End': eng['End'],
                'English Subtitle': eng['Text']
            })

        df = pd.DataFrame(data)
        base_name = os.path.splitext(os.path.basename(english_path))[0]
        if file_extension_en.lower()=='.srt':
            excel_file_name = f"{base_name}{tail}_srt.xlsx"
        elif file_extension_en.lower()=='.vtt':
            excel_file_name = f"{base_name}{tail}_vtt.xlsx"

    elif japanese_path:
        _, file_extension_ja = os.path.splitext(japanese_path)
        if file_extension_ja.lower()=='.srt':
            japanese_subtitles = parse_srt(japanese_path)
        elif file_extension_ja.lower()=='.vtt':
            japanese_subtitles = parse_vtt(japanese_path)
        else:
            print('入力に誤りがあります。')
            return None,None
        data = []
        for jap in japanese_subtitles:
            data.append({
                'ID': jap['ID'],
                'Start': jap['Start'],
                'End': jap['End'],
                f'{lang_name} Subtitle': jap['Text']
            })
        
        

        df = pd.DataFrame(data)
        base_name = os.path.splitext(os.path.basename(japanese_path))[0]
        if file_extension_ja.lower()=='.srt':
            excel_file_name = f"{base_name}{tail}_srt.xlsx"
        elif file_extension_ja.lower()=='.vtt':
            excel_file_name = f"{base_name}{tail}_vtt.xlsx"
    else:
        return None, None

    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    temp_dir = os.path.join(tempfile.gettempdir(), f"tempdir_{timestamp}")
    os.makedirs(temp_dir, exist_ok=True)
    temp_file_path = os.path.join(temp_dir, excel_file_name)
   
   
    
    with pd.ExcelWriter(temp_file_path, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Subtitles')
        workbook = writer.book
        worksheet = writer.sheets['Subtitles']

        column_widths = {'A': 7, 'B': 25, 'C': 25, 'D': 90, 'E': 90}
        for column, width in column_widths.items():
            worksheet.column_dimensions[column].width = width

        for row in worksheet.iter_rows(min_row=2, max_row=len(df) + 1):
            for cell in row:
                if cell.column_letter == 'A':
                    cell.alignment = Alignment(horizontal='right', vertical='center')
                elif cell.column_letter in ['B', 'C']:
                    cell.alignment = Alignment(horizontal='center', vertical='center')
                elif cell.column_letter in ['D', 'E']:
                    cell.alignment = Alignment(horizontal='left', vertical='center')

        for row in worksheet.iter_rows(min_row=2, max_row=len(df) + 1):
            worksheet.row_dimensions[row[0].row].height = 30

        header_font = Font(bold=True)
        for cell in worksheet["1:1"]:
            cell.font = header_font
            cell.alignment = Alignment(horizontal='center')
            cell.fill = PatternFill(start_color="DAEEF3", end_color="DAEEF3", fill_type="solid")
        


    return temp_file_path, df



#tab5
import re
from openpyxl.styles import Alignment, Font, PatternFill
import pandas as pd


def format_timestamp(seconds):
    hrs, secs = divmod(seconds, 3600)
    mins, secs = divmod(secs, 60)
    millis = int((secs % 1) * 1000)
    return f"{int(hrs):02}:{int(mins):02}:{int(secs):02},{millis:03}"

#dataframe追加
def parse_srt_c(srt_content):
    pattern = r'(\d+)\n(\d{2}:\d{2}:\d{2},\d{3}) --> (\d{2}:\d{2}:\d{2},\d{3})\n(.*?)\n\n'
    matches = re.findall(pattern, srt_content, re.DOTALL)
    
    subtitles = []
    for match in matches:
        subtitles.append({
            'ID': int(match[0]),
            'Start': match[1],
            'End': match[2],
            'Text': match[3].replace('\n', ' ')
        })
    
    return subtitles

def dataframe_to_html_table(df):
    return df.to_html(index=False)

# SRTファイルからExcelファイルを作成する関数
def create_excel_from_srt_c(srt_content, input_file_name):
    excel_file_name = f"{input_file_name}_srt.xlsx"
    english_subtitles = parse_srt_c(srt_content)

    data = []
    for eng in english_subtitles:
        data.append({
            'ID': eng['ID'],
            'Start': eng['Start'],
            'End': eng['End'],
            'English Subtitle': eng['Text']
        })

    df = pd.DataFrame(data)

    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    temp_dir = os.path.join(tempfile.gettempdir(), f"tempdir_{timestamp}")
    os.makedirs(temp_dir, exist_ok=True)

    excel_file_path = os.path.join(temp_dir, excel_file_name)
    
    with pd.ExcelWriter(excel_file_path, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Subtitles')
        workbook = writer.book
        worksheet = writer.sheets['Subtitles']

        column_widths = {'A': 7, 'B': 25, 'C': 25, 'D': 90, 'E': 90}
        for column, width in column_widths.items():
            worksheet.column_dimensions[column].width = width

        for row in worksheet.iter_rows(min_row=2, max_row=len(df) + 1):
            for cell in row:
                if cell.column_letter == 'A':
                    cell.alignment = Alignment(horizontal='right', vertical='center')
                elif cell.column_letter in ['B', 'C']:
                    cell.alignment = Alignment(horizontal='center', vertical='center')
                elif cell.column_letter in ['D', 'E']:
                    cell.alignment = Alignment(horizontal='left', vertical='center')

        for row in worksheet.iter_rows(min_row=2, max_row=len(df) + 1):
            worksheet.row_dimensions[row[0].row].height = 30

        header_font = Font(bold=True)
        for cell in worksheet["1:1"]:
            cell.font = header_font
            cell.alignment = Alignment(horizontal='center')
            cell.fill = PatternFill(start_color="DAEEF3", end_color="DAEEF3", fill_type="solid")

    
    return excel_file_path, df


##tab6
def new_process_file(input_files,replace_word):
    '''if input_file is None:
            return None, None, [],pd.DataFrame({'1': [''], '2': [''], '3': ['']})#0729
    '''
    output_srts=[]
    output_excels=[]
    output_txts=[]

    for input_file in input_files:
        with open(input_file, 'r', encoding='utf-8') as file:
            lines = file.readlines()

        basename = os.path.splitext(os.path.basename(input_file))[0]
        
        timestamp_patch = datetime.now().strftime("%Y%m%d%H%M%S")
        temp_dir = os.path.join(tempfile.gettempdir(), f"tempdir_{timestamp_patch}")
        os.makedirs(temp_dir, exist_ok=True) 
        if input_file.lower().endswith('.vtt'):
            output = process_vtt(lines,replace_word)
            output_file = os.path.join(temp_dir, f"{basename}_ed.vtt")
            output_srts.append(output_file)
        elif input_file.lower().endswith('.srt'):
            output = process_srt(lines,replace_word)
            output_file = os.path.join(temp_dir, f"{basename}_ed.srt")
            output_srts.append(output_file)
        else:
            raise ValueError('Unsupported file format')

        with open(output_file, 'w', encoding='utf-8') as file:
            file.write(output)



        if input_file.endswith(".srt"):
            segment_pattern = r'(\d+)\n(\d{1,2}:\d{2}:\d{2},\d{3}\s*-->\s*\d{1,2}:\d{2}:\d{2},\d{3})\n(.*?)(?=\n\d+\n|\Z)'
            segments = re.findall(segment_pattern, output, re.DOTALL)  # 全てのセグメントを抽出
                        
            result = []

            # 各セグメントごとにID、タイムスタンプ、テキストを処理
            for segment in segments:
                segment_id, timestamp, text = segment
                text=text.replace("\t","").replace("\n"," ")
                result.append(text)

            final_texts=''.join(result) 
        
        elif input_file.endswith(".vtt"):

            #texts=moz_t7.webvtt_remover(texts)
            segment_pattern = r'(\d+)\n(\d{1,2}:\d{2}:\d{2}.\d{3}\s*-->\s*\d{1,2}:\d{2}:\d{2}.\d{3})\n(.*?)(?=\n\d+\n|\Z)'
            segments = re.findall(segment_pattern, output, re.DOTALL)  # 全てのセグメントを抽出
            
            result = []

            # 各セグメントごとにID、タイムスタンプ、テキストを処理
            for segment in segments:
                segment_id, timestamp, text = segment
                text=text.replace("\t","").replace("\n"," ")
                result.append(text)   
            final_texts=''.join(result) 


        '''        doc = Document()
        doc.add_paragraph(final_texts)'''
        
        timestamp_patch = datetime.now().strftime("%Y%m%d%H%M%S")
        temp_dir = os.path.join(tempfile.gettempdir(), f"tempdir_{timestamp_patch}")
        os.makedirs(temp_dir, exist_ok=True) 

        if input_file.lower().endswith('.vtt'):
            texts_file = os.path.join(temp_dir, f"{basename}_ed_NR.txt")
            output_txts.append(texts_file)
        elif input_file.lower().endswith('.srt'):
            texts_file = os.path.join(temp_dir, f"{basename}_ed_NR.txt")
            output_txts.append(texts_file)
        #doc.save(docx_file)
        with open(texts_file,"w",encoding="utf-8") as f:
            f.write(final_texts)

        t7_excel_file,t7_df=create_excel_from_srt(english_path=output_file,tail="")#0729
        output_excels.append(t7_excel_file)
        #output_html = f"""<head><meta charset="UTF-8"></head><body><pre style="white-space: pre-wrap; overflow-y: auto; height: 300px; word-wrap: break-word; padding: 10px; font-family: inherit; font-size: inherit;">{output}</pre></body>"""

    if len(output_excels)>1:
        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
        temp_dir = os.path.join(tempfile.gettempdir(), f"tempdir_{timestamp}")
        os.makedirs(temp_dir, exist_ok=True)
        zipped_excels_path = os.path.join(temp_dir, "edited_excels.zip")
        with zipfile.ZipFile(zipped_excels_path, 'w') as zipf:
            for file in output_excels:
                zipf.write(file, os.path.basename(file))
        zipped_txts_path = os.path.join(temp_dir, "edited_txts.zip")
        with zipfile.ZipFile(zipped_txts_path, 'w') as zipf:
            for file in output_txts:
                zipf.write(file, os.path.basename(file))
        zipped_srts_path = os.path.join(temp_dir, "edited_srts.zip")
        with zipfile.ZipFile(zipped_srts_path, 'w') as zipf:
            for file in output_srts:
                zipf.write(file, os.path.basename(file))
            
        return [zipped_excels_path],[zipped_txts_path],[zipped_srts_path]
            
            


    # return output_html, [output_file, docx_file], output_file
    return output_excels,output_txts,output_srts